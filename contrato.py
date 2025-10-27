from docx import Document
from datetime import datetime
documento = Document("contrato.docx")

# for paragrafo in documento.paragraphs:
#     print(paragrafo.text)
#     paragrafo.text = paragrafo.text.replace('XXXX', 'Bruno')

# documento.save("contrato_atualizado.docx")

nome = 'Bruno'
item1 = 'Serviço de Consultoria'
item2 = 'Desenvolvimento de Software'
item3 = 'Suporte Técnico'


referencias = {
    'XXXX': nome,
    'YYYY': item1,
    'ZZZZ': item2,
    'WWWW': item3,
    'DD': str(datetime.now().day),
    'MM': str(datetime.now().month),
    'AAAA': str(datetime.now().year)
}
for paragrafo in documento.paragraphs:
    for codigo in referencias:
        valor = referencias[codigo]
        paragrafo.replace(codigo, valor)
        

documento.save("contrato_atualizado.docx")

